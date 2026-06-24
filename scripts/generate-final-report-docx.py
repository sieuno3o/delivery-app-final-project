from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "docx" / "동네한입_기말프로젝트_최종보고서.docx"
TMP = ROOT / "tmp" / "docx-report"
DIAGRAM = TMP / "database-relationships.png"

# Preset: compact_reference_guide.
# Named override `a4_submission_compact`: A4 portrait, 15/16 mm margins,
# 8.7 pt Korean body text, 180 mm / 10,205 DXA usable width, 1.12 line spacing.
# Named override `dongne_hanip_brand`: #F04424 accent, #211F1C ink.
PAGE_WIDTH_DXA = 11906
CONTENT_WIDTH_DXA = 10205
TABLE_WIDTH_DXA = 10200
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 70, "bottom": 70, "start": 120, "end": 120}

FONT = "Apple SD Gothic Neo"
FONT_PATH = "/System/Library/Fonts/AppleSDGothicNeo.ttc"
INK = "211F1C"
ACCENT = "F04424"
MUTED = "6F6A63"
LINE = "DED9D0"
SURFACE = "F5F2EB"
SOFT_ORANGE = "FFF0E8"
SOFT_GREEN = "EAF7F0"
WHITE = "FFFFFF"


def set_run_font(run, size=None, color=INK, bold=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def remove_paragraph_borders(element):
    if hasattr(element, "_p"):
        properties = element._p.get_or_add_pPr()
    else:
        properties = element._element.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is not None:
        properties.remove(borders)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins=CELL_MARGINS_DXA):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in margins.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def apply_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color=LINE, size=4):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def mark_row_as_header(row):
    row_properties = row._tr.get_or_add_trPr()
    header = row_properties.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        row_properties.append(header)
    header.set(qn("w:val"), "1")


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), ACCENT)
    run_properties.append(color)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)
    run_properties.append(fonts)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    run_properties.append(underline)
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph):
    run = paragraph.add_run()
    set_run_font(run, 7.5, MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])


def create_numbering(document):
    numbering = document.part.numbering_part.element
    existing_abstract = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    existing_num = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(existing_abstract, default=0) + 1
    num_id = max(existing_num, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    paragraph_props = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "260")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "40")
    spacing.set(qn("w:line"), "264")
    spacing.set(qn("w:lineRule"), "auto")
    paragraph_props.extend([tabs, indent, spacing])
    level.extend([start, number_format, level_text, justification, paragraph_props])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_number(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.extend([level, number])
    p_pr.append(num_pr)


def draw_diagram(path):
    width, height = 2400, 900
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = ImageFont.truetype(FONT_PATH, 34)
    key_font = ImageFont.truetype(FONT_PATH, 27)
    small = ImageFont.truetype(FONT_PATH, 23)
    stroke = "#DED9D0"
    accent = "#F04424"
    fill = "#F5F2EB"
    fill_order = "#FFF0E8"
    ink = "#211F1C"
    muted = "#6F6A63"

    positions = {
        "users": (40, 90, 500, 350),
        "sessions": (40, 560, 500, 820),
        "restaurants": (620, 90, 1100, 350),
        "menu_items": (620, 560, 1100, 820),
        "orders": (1220, 90, 1740, 370),
        "order_items": (1220, 540, 1740, 840),
        "order_status_history": (1860, 90, 2360, 390),
    }

    def edge(points, label, label_at):
        draw.line(points, fill=stroke, width=6, joint="curve")
        end_x, end_y = points[-1]
        draw.ellipse((end_x - 7, end_y - 7, end_x + 7, end_y + 7), fill=accent)
        lx, ly = label_at
        bbox = draw.textbbox((0, 0), label, font=small)
        label_width = bbox[2] - bbox[0]
        label_height = bbox[3] - bbox[1]
        draw.rounded_rectangle(
            (lx - 12, ly - 8, lx + label_width + 12, ly + label_height + 8),
            radius=12,
            fill="white",
            outline=stroke,
            width=2,
        )
        draw.text((lx, ly), label, font=small, fill=muted)

    # Draw relationships behind the table boxes.
    edge([(270, 350), (270, 560)], "1:N", (292, 430))
    edge([(860, 350), (860, 560)], "1:N", (882, 430))
    edge([(1480, 370), (1480, 540)], "1:N", (1502, 430))
    edge([(1100, 220), (1220, 220)], "1:N", (1118, 175))
    edge([(1740, 225), (1860, 225)], "1:N", (1758, 180))
    edge([(1100, 690), (1220, 690)], "1:N", (1118, 645))
    edge([(500, 160), (540, 35), (1180, 35), (1220, 160)], "1:N", (820, 4))

    rows = {
        "users": ["PK  id", "UK  email", "role · password_hash"],
        "sessions": ["PK  id", "FK  user_id → users.id", "token_hash · expires_at"],
        "restaurants": ["PK  id", "UK  slug", "delivery_fee · minimum_order"],
        "menu_items": ["PK  id", "FK  restaurant_id → restaurants.id", "price · is_sold_out"],
        "orders": [
            "PK  id",
            "FK  user_id → users.id",
            "FK  restaurant_id → restaurants.id",
            "status · total · idempotency_key",
        ],
        "order_items": [
            "PK  id",
            "FK  order_id → orders.id",
            "FK  menu_item_id → menu_items.id",
            "menu_name · unit_price · quantity",
        ],
        "order_status_history": [
            "PK  id",
            "FK  order_id → orders.id",
            "FK  changed_by_user_id → users.id",
            "status · note · created_at",
        ],
    }

    for name, box in positions.items():
        x1, y1, x2, y2 = box
        box_fill = fill_order if name in {"orders", "order_items"} else fill
        box_stroke = accent if name in {"orders", "order_items"} else stroke
        draw.rounded_rectangle(box, radius=22, fill=box_fill, outline=box_stroke, width=5)
        header_bottom = y1 + 62
        draw.rounded_rectangle(
            (x1, y1, x2, header_bottom + 18),
            radius=22,
            fill=accent if name in {"orders", "order_items"} else ink,
        )
        draw.rectangle((x1, header_bottom, x2, header_bottom + 18), fill=accent if name in {"orders", "order_items"} else ink)
        bbox = draw.textbbox((0, 0), name, font=title_font)
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
        draw.text(
            ((x1 + x2 - tw) / 2, y1 + (62 - th) / 2 - 3),
            name,
            font=title_font,
            fill="white",
        )
        row_y = y1 + 90
        for index, row_text in enumerate(rows[name]):
            row_color = ink if index < 3 else muted
            draw.text((x1 + 22, row_y), row_text, font=key_font, fill=row_color)
            row_y += 46

    legend = "PK 기본키   ·   FK 외래키   ·   UK 고유키   ·   1:N 일대다 관계"
    legend_bbox = draw.textbbox((0, 0), legend, font=small)
    legend_width = legend_bbox[2] - legend_bbox[0]
    draw.text(((width - legend_width) / 2, 860), legend, font=small, fill=muted)
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, format="PNG", optimize=True)


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(8.7)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2.5)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in [
        ("Title", 22, INK, 0, 3),
        ("Subtitle", 9, MUTED, 0, 5),
        ("Heading 1", 13.5, ACCENT, 7, 3),
        ("Heading 2", 10.5, INK, 4.5, 1.5),
        ("Heading 3", 9.5, INK, 3, 1),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        if name == "Title":
            remove_paragraph_borders(style)

    for style_name in ["Caption", "Header", "Footer"]:
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(7.5)
        style.font.color.rgb = RGBColor.from_string(MUTED)


def add_label_detail(paragraph, label, detail):
    label_run = paragraph.add_run(f"{label}  ")
    set_run_font(label_run, 8.5, ACCENT, bold=True)
    detail_run = paragraph.add_run(detail)
    set_run_font(detail_run, 8.5, INK)


def add_bug_section(
    document,
    number,
    title,
    symptom,
    cause,
    fix,
    video=False,
    investigation=None,
):
    heading = document.add_heading(level=2)
    prefix = heading.add_run(f"{number}  ")
    set_run_font(prefix, 10.5, ACCENT, bold=True)
    title_run = heading.add_run(title)
    set_run_font(title_run, 10.5, INK, bold=True)
    if video:
        badge = heading.add_run("   [영상 설명 버그]")
        set_run_font(badge, 8, ACCENT, bold=True)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    details = [("증상·영향", symptom)]
    if investigation:
        details.append(("첫 추측·측정", investigation))
    details.extend([("실제 원인", cause), ("해결·검증", fix)])
    for index, (label, text) in enumerate(details):
        label_run = paragraph.add_run(f"{label}  ")
        set_run_font(label_run, 8.2, MUTED, bold=True)
        text_run = paragraph.add_run(text)
        set_run_font(text_run, 8.2, INK)
        if index < len(details) - 1:
            text_run.add_break()


def add_callout(document, text, fill=SOFT_ORANGE, border=ACCENT):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(7)
    paragraph.paragraph_format.right_indent = Pt(7)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.12
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "5")
        tag.set(qn("w:space"), "5")
        tag.set(qn("w:color"), border)
        borders.append(tag)
    properties.append(borders)
    run = paragraph.add_run(text)
    set_run_font(run, 7.9, INK)


def make_document():
    document = Document()
    document.core_properties.title = "동네한입 기말 프로젝트 최종 보고서"
    document.core_properties.author = "김시은"
    document.core_properties.subject = "배달앱 데이터베이스 구조, 실제 버그 해결, AI 한계"
    document.core_properties.keywords = "Next.js, PostgreSQL, Docker, Neon, Vercel"

    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(14)
    section.bottom_margin = Mm(12)
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)
    section.header_distance = Mm(6)
    section.footer_distance = Mm(6)

    configure_styles(document)
    num_id = create_numbering(document)

    header = section.header.paragraphs[0]
    header.style = document.styles["Header"]
    header.paragraph_format.tab_stops.add_tab_stop(Mm(180))
    left = header.add_run("컴퓨터과학개론 · 111873-101")
    set_run_font(left, 7.3, MUTED, bold=True)
    right = header.add_run("\t기말 프로젝트 - 배달앱 만들기 & 배포")
    set_run_font(right, 7.3, MUTED)

    footer = section.footer.paragraphs[0]
    footer.style = document.styles["Footer"]
    footer.paragraph_format.tab_stops.add_tab_stop(Mm(180))
    footer.add_run("동네한입 · 김시은\t")
    for run in footer.runs:
        set_run_font(run, 7.3, MUTED)
    add_page_field(footer)
    suffix = footer.add_run(" / 2")
    set_run_font(suffix, 7.3, MUTED)

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(1)
    run = kicker.add_run("FINAL PROJECT · DELIVERY APP")
    set_run_font(run, 8.3, ACCENT, bold=True)

    title = document.add_paragraph(style="Title")
    remove_paragraph_borders(title)
    run = title.add_run("동네한입 - 기말 프로젝트 최종 보고서")
    set_run_font(run, 22, INK, bold=True)

    metadata = document.add_paragraph(style="Subtitle")
    metadata.paragraph_format.space_after = Pt(2)
    run = metadata.add_run("작성자 김시은  ·  작성일 2026.06.24  ·  Next.js / PostgreSQL / Docker / Vercel")
    set_run_font(run, 8.2, MUTED)

    links = document.add_paragraph()
    links.paragraph_format.space_after = Pt(4)
    label = links.add_run("운영 URL  ")
    set_run_font(label, 7.8, MUTED, bold=True)
    add_hyperlink(links, "dongne-hanip-kohl.vercel.app", "https://dongne-hanip-kohl.vercel.app")
    separator = links.add_run("    GitHub  ")
    set_run_font(separator, 7.8, MUTED, bold=True)
    add_hyperlink(
        links,
        "sieuno3o/delivery-app-final-project",
        "https://github.com/sieuno3o/delivery-app-final-project",
    )

    document.add_heading("1. 데이터베이스 구조와 설계 이유", level=1)
    picture_paragraph = document.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.space_after = Pt(1)
    picture = picture_paragraph.add_run().add_picture(str(DIAGRAM), width=Mm(176))
    picture._inline.docPr.set(
        "descr",
        "7개 테이블의 PK, 주요 FK와 일대다 관계를 표시한 데이터베이스 구조도",
    )
    caption = document.add_paragraph("그림 1. 핵심 PK/FK와 1:N 관계 - 상세 컬럼은 아래 표와 저장 흐름에서 설명한다.", style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(3)

    rows = [
        ("users", "계정, bcrypt 비밀번호 해시, 역할", "PK users.id"),
        ("sessions", "세션 토큰 해시와 만료 시각", "PK sessions.id\nFK user_id → users.id"),
        ("restaurants", "식당, 배달비, 최소 주문 금액", "PK restaurants.id"),
        ("menu_items", "식당별 메뉴, 가격, 품절 여부", "PK menu_items.id\nFK restaurant_id → restaurants.id"),
        ("orders", "주문자·식당·배송지·상태·총액", "PK orders.id\nFK user_id → users.id\nFK restaurant_id → restaurants.id"),
        ("order_items", "메뉴명·단가 스냅샷과 수량", "PK order_items.id\nFK order_id → orders.id\nFK menu_item_id → menu_items.id"),
        ("order_status_history", "상태 변경 시각·변경자·메모", "PK order_status_history.id\nFK order_id → orders.id\nFK changed_by_user_id → users.id"),
    ]
    table = document.add_table(rows=1, cols=3)
    apply_table_geometry(table, [2200, 3500, 4500])
    set_table_borders(table)
    mark_row_as_header(table.rows[0])
    headers = table.rows[0].cells
    for index, text in enumerate(["테이블", "한 줄 역할", "PK / 주요 FK"]):
        set_cell_shading(headers[index], INK)
        paragraph = headers[index].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, 7.4, WHITE, bold=True)
    for name, role, keys in rows:
        cells = table.add_row().cells
        for index, text in enumerate([name, role, keys]):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(text)
            set_run_font(run, 6.8, INK, bold=index == 0)

    document.add_heading("왜 이렇게 나눴는가", level=2)
    paragraph = document.add_paragraph()
    add_label_detail(
        paragraph,
        "주문 ↔ 주문상세",
        "한 주문에 여러 메뉴가 들어가므로 공통값은 orders, 반복값은 order_items로 분리했다. 메뉴명·단가 스냅샷을 남겨 메뉴 수정·삭제 후에도 과거 결제 기록을 보존한다.",
    )
    paragraph = document.add_paragraph()
    add_label_detail(
        paragraph,
        "현재 상태 ↔ 상태 이력",
        "빠른 목록 조회는 orders.status, 변경 과정은 order_status_history에 누적한다. 조회 성능과 추적 가능성을 함께 얻는 의도적인 중복이다.",
    )
    paragraph = document.add_paragraph()
    add_label_detail(
        paragraph,
        "사용자 ↔ 세션",
        "한 사용자의 여러 기기 로그인을 허용하고 세션별 만료를 관리한다. DB에는 탈취 위험을 줄이기 위해 토큰 원문 대신 SHA-256 해시만 저장한다.",
    )

    document.add_heading("주문 저장 순서·일관성·권한", level=2)
    steps = [
        "로그인 사용자와 권한을 확인한다.",
        "식당·메뉴를 DB에서 다시 읽어 현재 가격, 품절, 최소 주문 금액을 검증한다.",
        "클라이언트 합계를 신뢰하지 않고 상품 합계·배달비·최종 금액을 서버에서 재계산한다.",
        "한 트랜잭션에서 orders 1행, order_items 여러 행, 최초 상태 이력을 저장한다.",
        "모두 성공하면 커밋하고 하나라도 실패하면 롤백한다. UUID 고유키로 중복 주문을 막고, 조회에는 현재 user_id를 함께 사용한다.",
    ]
    for text in steps:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1.5)
        apply_number(paragraph, num_id)
        run = paragraph.add_run(text)
        set_run_font(run, 8.1, INK)

    add_callout(
        document,
        "환경 분리  |  로컬은 Docker Compose의 PostgreSQL 16(호스트 5433), 운영은 Neon PostgreSQL을 사용한다. 두 환경에 같은 Drizzle 마이그레이션을 적용해 스키마 차이로 인한 배포 오류를 줄였다.",
    )

    kicker = document.add_paragraph()
    kicker.paragraph_format.page_break_before = True
    kicker.paragraph_format.space_after = Pt(1)
    run = kicker.add_run("ACTUAL PROBLEMS · VERIFICATION · LIMITS")
    set_run_font(run, 8.3, ACCENT, bold=True)
    document.add_heading("2. 직접 겪고 해결한 문제 3개", level=1)

    add_bug_section(
        document,
        "BUG-004",
        "검은 링크 버튼의 글자가 보이지 않음",
        "회원가입·식당 이동 같은 핵심 버튼의 텍스트를 알아볼 수 없었다. 콘솔 오류는 없었다.",
        "레이어 밖의 전역 a { color: inherit }가 Tailwind @layer utilities 안의 text-white보다 캐스케이드에서 우선했다. 선택자 명시도 문제가 아니라 cascade layer 순서 문제였다.",
        "불필요한 전역 a 규칙을 제거했다(커밋 e62c1e3). getComputedStyle 글자색이 rgb(255,255,255)로 바뀌고 실제 화면에서도 글자가 보이는지 확인했다.",
        video=True,
        investigation="Tailwind 클래스 생성 또는 캡처 문제라고 추측했다. 그러나 getComputedStyle($0).color로 측정하니 배경과 글자가 모두 rgb(33,31,28)이어서 CSS 덮어쓰기로 범위를 좁혔다.",
    )
    add_bug_section(
        document,
        "BUG-005",
        "주문 완료 안내가 즉시 사라짐",
        "주문은 DB에 저장됐지만 ‘주문이 접수됐어요!’ 안내가 거의 즉시 일반 주문 상세로 바뀌었다.",
        "장바구니를 비운 클라이언트 효과가 router.replace를 실행해 서버 컴포넌트를 재렌더링하고, 완료 화면 조건인 placed=1을 즉시 제거했다.",
        "history.replaceState로 주소만 정리했다. 현재 완료 화면은 유지되고 새로고침할 때만 일반 상세가 열리는지 실제 주문으로 검증했다.",
        investigation="서버 redirect가 placed=1을 전달하지 못했다고 생각했다. 네트워크 기록에서 완료 화면 뒤 주문 상세가 다시 요청되는 것을 확인해 클라이언트 효과를 추적했다.",
    )
    add_bug_section(
        document,
        "BUG-006",
        "운영 주문 시간이 한국 시간보다 9시간 느림",
        "같은 주문이 로컬에서는 한국 시간, Vercel에서는 정확히 9시간 전으로 표시됐다.",
        "PostgreSQL timestamptz의 UTC 저장은 정상이었지만 Intl.DateTimeFormat이 실행 서버의 기본 시간대에 의존했다. 로컬 macOS는 KST, Vercel 서버는 UTC였다.",
        "주문 날짜 포맷에 timeZone: Asia/Seoul을 명시하고 운영 재배포 후 동일 주문의 한국 시간 표시를 확인했다.",
        investigation="Neon의 UTC 저장 오류를 먼저 의심했다. 그러나 DB 시각은 정상이고 정확히 9시간 차이가 반복돼 저장이 아닌 표시 환경 문제로 범위를 좁혔다.",
    )

    document.add_heading("3. AI 활용의 한계와 남은 범위", level=1)
    paragraph = document.add_paragraph()
    add_label_detail(
        paragraph,
        "AI로 바로 안 풀린 것",
        "CSS는 명시도 추측만으로 해결되지 않아 computed style·CSS layer·Git diff를 비교했고, 시간 오류도 DB 값과 실행 환경을 따로 확인했다.",
    )
    paragraph = document.add_paragraph()
    add_label_detail(
        paragraph,
        "검증 방식",
        "DB·인증·장바구니·주문·배포·버그·테스트를 기능 단위로 커밋하고 실제 브라우저·DB·자동 검사로 다시 확인했다.",
    )
    paragraph = document.add_paragraph()
    add_label_detail(
        paragraph,
        "남은 범위",
        "실제 결제·이메일 인증·실시간 라이더 위치는 외부 연동의 보안·비용 때문에 제외했고, 물리 기기 확인·영상 녹화·제출은 직접 수행한다.",
    )

    document.add_heading("4. 검증 결과", level=1)
    verification = [
        ("필수 기능", "공개 URL: 회원가입 · 로그인 · 로그아웃 · 식당/메뉴 · 장바구니 · 주문 저장 · 내 주문내역", "7 / 7"),
        ("추가 기능", "검색·복합 필터 · 품절 · 상태 타임라인 · 관리자 변경 · 모바일 UI", "구현"),
        ("6/24 재검증", "ESLint · TypeScript · Vitest 28개 · Next.js 프로덕션 빌드", "통과"),
        ("Chrome E2E", "가입 → 주문 → 관리자 변경 → 고객 확인, 필터, 360px", "3 / 3"),
        ("로컬 DB", "Docker PostgreSQL · 식당 6개/메뉴 30개 · 동일 마이그레이션", "통과"),
        ("운영", "Vercel 공개 URL과 /login 경로 HTTP 200 (2026-06-24)", "통과"),
    ]
    table = document.add_table(rows=1, cols=3)
    apply_table_geometry(table, [1950, 6850, 1400])
    set_table_borders(table, "C8E6D4", 4)
    mark_row_as_header(table.rows[0])
    for index, text in enumerate(["영역", "검증 내용", "결과"]):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, SOFT_GREEN)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, 7.3, INK, bold=True)
    for area, evidence, result in verification:
        cells = table.add_row().cells
        for index, text in enumerate([area, evidence, result]):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if index in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
            )
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(text)
            set_run_font(run, 7.1, INK, bold=index == 2)

    add_callout(
        document,
        "제출물 일치  |  영상 상세 버그는 보고서의 BUG-004와 같고, 영상·보고서·README의 운영 URL, GitHub 주소, 7개 DB 테이블과 필수 기능 목록을 동일하게 맞췄다.",
        fill=SURFACE,
        border=LINE,
    )

    add_callout(
        document,
        "핵심 배움  |  화면 기능을 늘리는 것보다 데이터가 왜 나뉘는지, 실패할 때 트랜잭션·제약·권한 검사로 어떻게 일관성을 지키는지가 더 중요했다.",
        fill=SOFT_ORANGE,
        border=ACCENT,
    )

    return document


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    TMP.mkdir(parents=True, exist_ok=True)
    draw_diagram(DIAGRAM)
    document = make_document()
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
